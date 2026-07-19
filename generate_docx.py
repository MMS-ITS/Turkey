#!/usr/bin/env python3
import json, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

records = json.load(open('records_final.json', encoding='utf-8'))

# ---- Sort alphabetically by primary company name ----
def sort_key(r):
    if not r['companies']: return ('ZZZ',)
    n = r['companies'][0][0].upper()
    n = re.sub(r'^[^A-Z0-9ÇĞİÖŞÜ]+', '', n)
    return (n,)
records.sort(key=sort_key)

# ---- Colours ----
NAVY   = '1F3864'   # header background
LIGHT  = 'DCE6F1'   # banded row
WHITE  = 'FFFFFF'
ACCENT = '2E74B5'   # links / accents

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag,val in (('top',top),('bottom',bottom),('start',left),('end',right)):
        e = OxmlElement('w:'+tag); e.set(qn('w:w'), str(val)); e.set(qn('w:type'),'dxa'); m.append(e)
    tcPr.append(m)

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader'); th.set(qn('w:val'),'true'); trPr.append(th)

def set_table_borders(table, color='B4C6E7', sz=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement('w:'+edge)
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),str(sz))
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),color)
        borders.append(e)
    tblPr.append(borders)

def no_space(par):
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    par.paragraph_format.line_spacing = 1.0

# ---- Document setup (landscape for wide table) ----
doc = Document()
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.left_margin = sec.right_margin = Cm(1.3)
sec.top_margin = sec.bottom_margin = Cm(1.3)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9)

# ---- Title block ----
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('Yarn Manufacturer Exhibitor Directory')
run.bold = True; run.font.size = Pt(20); run.font.color.rgb = RGBColor.from_string(NAVY)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Consolidated & Verified Contact List')
r2.font.size = Pt(11); r2.italic = True; r2.font.color.rgb = RGBColor.from_string(ACCENT)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
rm = meta.add_run(f'{len(records)} entries  \u2022  email domains verified via live DNS (MX/A)  \u2022  missing emails sourced from company websites')
rm.font.size = Pt(8); rm.font.color.rgb = RGBColor.from_string('808080')
doc.add_paragraph()

# ---- Table ----
table = doc.add_table(rows=1, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.allow_autofit = False
set_table_borders(table)

widths = [Cm(1.0), Cm(7.6), Cm(6.2), Cm(6.6), Cm(5.2)]
headers = ['No.', 'Name of Company and Country', 'Website', 'Email', 'Products']

hdr = table.rows[0]
set_repeat_header(hdr)
for i, cell in enumerate(hdr.cells):
    shade(cell, NAVY)
    set_cell_margins(cell)
    cell.width = widths[i]
    p = cell.paragraphs[0]; no_space(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i in (1,) else WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(headers[i])
    run.bold = True; run.font.color.rgb = RGBColor.from_string(WHITE); run.font.size = Pt(10)

def add_lines(cell, items, color=None, size=8.5, sep_blank=False):
    """add each item as its own paragraph"""
    first = True
    for it in items:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        no_space(p); first = False
        run = p.add_run(it)
        run.font.size = Pt(size)
        if color: run.font.color.rgb = RGBColor.from_string(color)
    if first:  # nothing added -> keep empty paragraph
        no_space(cell.paragraphs[0])

n = 0
for r in records:
    n += 1
    row = table.add_row()
    banded = (n % 2 == 0)
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]
        set_cell_margins(cell)
        if banded: shade(cell, LIGHT)

    # No.
    p = row.cells[0].paragraphs[0]; no_space(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(str(n)).font.size = Pt(8.5)

    # Company + country (company bold)
    ccell = row.cells[1]; first = True
    for (name, country) in r['companies']:
        p = ccell.paragraphs[0] if first else ccell.add_paragraph()
        no_space(p); first = False
        rn = p.add_run(name); rn.bold = True; rn.font.size = Pt(9)
        rn.font.color.rgb = RGBColor.from_string(NAVY)
        if country:
            pc = ccell.add_paragraph(); no_space(pc)
            rc = pc.add_run(country); rc.italic = True; rc.font.size = Pt(8)
            rc.font.color.rgb = RGBColor.from_string('595959')

    # Website(s)
    add_lines(row.cells[2], r['websites'], color=ACCENT, size=8.5)

    # Email(s)
    add_lines(row.cells[3], r['emails'], color='000000', size=8.5)

    # Products (brief, comma-joined)
    pcell = row.cells[4]; pp = pcell.paragraphs[0]; no_space(pp)
    prods = ', '.join(r['products']) if r['products'] else '\u2014'
    rp = pp.add_run(prods); rp.font.size = Pt(8.5); rp.font.color.rgb = RGBColor.from_string('333333')

# ---- Summary section ----
only_web = sum(1 for r in records if r['websites'] and not r['emails'])
only_mail = sum(1 for r in records if r['emails'] and not r['websites'])
both = sum(1 for r in records if r['emails'] and r['websites'])
neither = sum(1 for r in records if not r['emails'] and not r['websites'])

doc.add_paragraph()
h = doc.add_paragraph()
rh = h.add_run('Summary'); rh.bold = True; rh.font.size = Pt(13)
rh.font.color.rgb = RGBColor.from_string(NAVY)

stab = doc.add_table(rows=5, cols=2)
stab.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(stab, color='B4C6E7')
srows = [
    ('Entries containing only a website, but no email address', only_web),
    ('Entries containing only an email address, but no website', only_mail),
    ('Entries containing neither a website nor an email address', neither),
    ('Entries containing both a website and an email address', both),
    ('Total entries', len(records)),
]
for i,(label,val) in enumerate(srows):
    c0, c1 = stab.rows[i].cells
    c0.width = Cm(14); c1.width = Cm(3)
    set_cell_margins(c0); set_cell_margins(c1)
    is_total = (i==4)
    shade(c0, NAVY if is_total else LIGHT); shade(c1, NAVY if is_total else 'FFFFFF')
    p0 = c0.paragraphs[0]; no_space(p0)
    r0 = p0.add_run(label); r0.font.size = Pt(10); r0.bold = is_total
    r0.font.color.rgb = RGBColor.from_string(WHITE if is_total else '333333')
    p1 = c1.paragraphs[0]; no_space(p1); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(str(val)); r1.bold = True; r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor.from_string(WHITE if is_total else NAVY)

# footer note
doc.add_paragraph()
note = doc.add_paragraph()
rnote = note.add_run(
    'Notes: Duplicate entries and superfluous details were removed. Where one email address is shared by '
    'multiple companies, those companies are grouped in a single row. Multiple emails/websites for one company '
    'are combined in one cell, separated by line breaks. Email domains were validated with live DNS lookups; '
    'only domains with no MX and no A record were discarded. For companies listed with a website but no email, '
    'the website was fetched live to recover a published contact address where available.')
rnote.font.size = Pt(7.5); rnote.italic = True; rnote.font.color.rgb = RGBColor.from_string('808080')

out = 'Yarn Manufacturer Exhibitor List - Output.docx'
doc.save(out)
print('Saved', out)
print('rows:', len(records), '| only_web', only_web, '| only_mail', only_mail, '| both', both, '| neither', neither)
